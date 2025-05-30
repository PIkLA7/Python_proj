from flask import Flask, render_template, redirect, url_for, request, session, flash
from flask_mysqldb import MySQL
from functools import wraps
import MySQLdb
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
from io import BytesIO
from flask import send_file
from openpyxl import Workbook, load_workbook
import os
from datetime import datetime
from datetime import timedelta
from docx import Document
from docx.shared import Inches
from zipfile import ZipFile
import random
from openpyxl.styles import Font
from collections import defaultdict
from dotenv import load_dotenv

load_dotenv()  # Загружаем переменные из .env
app = Flask(__name__)

# Настройки подключения к базе данных
app.config['MYSQL_HOST'] = os.getenv('MYSQL_HOST', 'localhost')
app.config['MYSQL_USER'] = os.getenv('MYSQL_USER', 'root')
app.config['MYSQL_PASSWORD'] = os.getenv('MYSQL_PASSWORD', '')
app.config['MYSQL_DB'] = os.getenv('MYSQL_DB', 'college_schedule')
app.config['MYSQL_CURSORCLASS'] = os.getenv('MYSQL_CURSORCLASS', 'DictCursor')

app.config['SECRET_PHRASE_ADMIN'] = os.getenv('SECRET_PHRASE_ADMIN', 'admin_secret')
app.config['SECRET_PHRASE_TEACHER'] = os.getenv('SECRET_PHRASE_TEACHER', 'teacher_secret')

REPORTS_FOLDER = os.getenv('REPORTS_FOLDER', 'downloads')
EXCEL_FILE_PATH = os.getenv('EXCEL_FILE_PATH', 'user_actions_report.xlsx')

app.secret_key = os.getenv('SECRET_KEY', 'your_secret_key')

mysql = MySQL(app)







REQUIRED_ENV_VARS = [
    'MYSQL_HOST', 'MYSQL_USER', 'MYSQL_PASSWORD', 'MYSQL_DB',
    'MYSQL_CURSORCLASS', 'SECRET_KEY',
    'SECRET_PHRASE_ADMIN', 'SECRET_PHRASE_TEACHER'
]

def is_configured():
    return all(os.getenv(var) for var in REQUIRED_ENV_VARS)

@app.route('/setup', methods=['GET', 'POST'])
def setup():
    if is_configured():
        return redirect(url_for('index'))

    if request.method == 'POST':
        config = {
            'MYSQL_HOST': request.form['MYSQL_HOST'],
            'MYSQL_USER': request.form['MYSQL_USER'],
            'MYSQL_PASSWORD': request.form['MYSQL_PASSWORD'],
            'MYSQL_DB': request.form['MYSQL_DB'],
            'MYSQL_CURSORCLASS': request.form['MYSQL_CURSORCLASS'],
            'SECRET_KEY': request.form['SECRET_KEY'],
            'SECRET_PHRASE_ADMIN': request.form['SECRET_PHRASE_ADMIN'],
            'SECRET_PHRASE_TEACHER': request.form['SECRET_PHRASE_TEACHER'],
            'REPORTS_FOLDER': request.form.get('REPORTS_FOLDER', 'downloads'),
            'EXCEL_FILE_PATH': request.form.get('EXCEL_FILE_PATH', 'user_actions_report.xlsx')
        }

        # Сохраняем в .env
        with open('.env', 'w') as f:
            for key, value in config.items():
                f.write(f"{key}={value}\n")

        flash("Настройка завершена. Перезапустите приложение.")
        return redirect(url_for('setup'))

    return render_template('setup.html')  # HTML-форма с полями конфигурации

# Главная страница
@app.route('/')
def index():
    if not is_configured():
        return render_template('not_ready.html')  # страница-заглушка
    return render_template('index.html')


# Личный кабинет
@app.route('/cabinet')
def cabinet():
    if 'user_id' not in session:
        flash('Вы должны войти в систему для доступа к кабинету.')
        return redirect(url_for('login'))
    
    role = session.get('role', 'user')
    return render_template('cabinet.html', role=role)


# Декоратор для проверки роли
def role_required(role):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session or session.get('role') != role:
                flash('У вас нет доступа к этому разделу')
                return redirect(url_for('login'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# Вход пользователя
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        cur = mysql.connection.cursor()
        cur.execute("SELECT * FROM users WHERE username = %s", (username,))
        user = cur.fetchone()
        cur.close()
        if user and check_password_hash(user['password_hash'], password):
            session['user_id'] = user['id']
            session['username'] = user['username'] 
            session['role'] = user['role']

            log_user_action(session.get('user_id'), "Просмотрено расписание", request.method)
            return redirect(url_for('cabinet'))
        flash('Неверное имя пользователя или пароль')
    return render_template('login.html')


# Выход из системы
@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))




def log_user_action(user_id, action, request_type):
    if not user_id:
        return  # Не логируем действия незарегистрированных пользователей
    ip_address = request.remote_addr  # Получаем IP-адрес пользователя
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")  # Время действия

    # Добавляем запись в таблицу user_actions
    cur = mysql.connection.cursor()
    cur.execute("""
        INSERT INTO user_actions (user_id, action, timestamp, request_type, ip_address)
        VALUES (%s, %s, %s, %s, %s)
    """, (user_id, action, timestamp, request_type, ip_address))
    mysql.connection.commit()
    cur.close()



    # Получаем имя пользователя по user_id
    cur = mysql.connection.cursor()
    cur.execute("SELECT username FROM users WHERE id = %s", (user_id,))
    user = cur.fetchone()
    cur.close()
    # Формируем новый DataFrame с действием пользователя
    new_data = pd.DataFrame([{
           "Username": user['username'],
           "Action": action,
           "Timestamp": timestamp,
           "Request Type": request_type,
           "IP Address": ip_address
       }])

    print("DataFrame с новым действием:")
    print(new_data)
    # Если файл Excel существует, читаем его
    if os.path.exists(EXCEL_FILE_PATH):
        try:
           existing_df = pd.read_excel(EXCEL_FILE_PATH)
           print("Существующий DataFrame перед добавлением:")
           print(existing_df)
           # Объединяем данные
           df = pd.concat([existing_df, new_data], ignore_index=True)
        except Exception as e:
           print(f"Ошибка при чтении существующего файла: {e}")
           df = new_data
    else:
       # Если файла нет, используем только новые данные
       df = new_data
    print("Итоговый DataFrame перед записью:")
    print(df)
    # Записываем данные в файл Excel
    try:
       with pd.ExcelWriter(EXCEL_FILE_PATH, engine='openpyxl') as writer:
           df.to_excel(writer, index=False, sheet_name='User Actions')
       print("Данные успешно записаны в Excel файл.")
    except Exception as e:
       print(f"Ошибка при записи в Excel файл: {e}")    



@app.route('/download_report')
@role_required('admin')
def download_report():
    cur = mysql.connection.cursor()
    query = "SELECT u.username, ua.action, ua.timestamp, ua.request_type, ua.ip_address FROM user_actions ua JOIN users u ON ua.user_id = u.id"
    cur.execute(query)
    actions = cur.fetchall()
    cur.close()

    # Преобразуем в DataFrame для экспорта в Excel
    df = pd.DataFrame(actions, columns=["Username", "Action", "Timestamp", "Request Type", "IP Address"])

    # Сохраняем в Excel файл
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Actions Report')
    output.seek(0)

    # Отправляем файл пользователю с новым параметром download_name
    return send_file(output, download_name="user_actions_report.xlsx", as_attachment=True)

# Регистрация нового пользователя
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        teacher_fullname = request.form.get('teacher_fullname', '')
        username = request.form['username']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        special_user = 'special_user' in request.form
        secret_phrase = request.form.get('secret_phrase', '')

        if password != confirm_password:
            flash('Пароли не совпадают')
            return redirect(url_for('register'))

        role = 'user'  # Роль по умолчанию
        if special_user:
            if secret_phrase == app.config['SECRET_PHRASE_ADMIN']:
                role = 'admin'
            elif secret_phrase == app.config['SECRET_PHRASE_TEACHER']:
                role = 'teacher'
            else:
                flash('Неверная секретная фраза')
                return redirect(url_for('register'))

        # Хешируем пароль и сохраняем пользователя
        hashed_password = generate_password_hash(password)
        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO users (username, password_hash, role)
            VALUES (%s, %s, %s)
        """, (username, hashed_password, role))
        if role == 'teacher':
            if not teacher_fullname:
                flash("Для преподавателя обязательно указать ФИО.")
                return redirect(url_for('register'))
            cur.execute("INSERT INTO teachers (full_name, username) VALUES (%s, %s)", (teacher_fullname, username))

        mysql.connection.commit()
        cur.close()
        log_user_action(session.get('user_id'), f"Зарегистрирован новый пользователь: {username}", request.method)

        flash('Регистрация прошла успешно')
        return redirect(url_for('login'))
    return render_template('registration.html')



@app.route('/view_schedule')
def view_schedule():
    cur = mysql.connection.cursor()

    cur.execute("""
        SELECT s.id, g.group_name, sub.name AS subject_name, t.full_name AS teacher_name, 
               s.date, s.time_slot, s.is_half_lesson, s.is_empty
        FROM schedules s
        JOIN grooups g ON s.group_id = g.id
        JOIN subjects sub ON s.subject_id = sub.id
        JOIN teachers t ON s.teacher_id = t.id
    """)
    schedule = cur.fetchall()

    # Загружаем замены
    cur.execute("""
        SELECT r.schedule_id, sub.name AS new_subject_name, t.full_name AS new_teacher_name
        FROM replacements r
        JOIN subjects sub ON r.new_subject_id = sub.id
        JOIN teachers t ON r.new_teacher_id = t.id
    """)
    replacements = {row['schedule_id']: row for row in cur.fetchall()}
    cur.close()

    # Подставляем замену
    for item in schedule:
        if item['id'] in replacements:
            item['subject_name'] = replacements[item['id']]['new_subject_name']
            item['teacher_name'] = replacements[item['id']]['new_teacher_name']

    role = session.get('role', 'guest')
    return render_template('view_schedule.html', schedule=schedule, role=role)


# Только для преподавателя: просмотр своих пар
@app.route('/teacher_schedule')
@role_required('teacher')
def teacher_schedule():
    username = session.get('username')
    cur = mysql.connection.cursor()
    cur.execute("SELECT id FROM teachers WHERE username = %s", (username,))
    teacher = cur.fetchone()
    if not teacher:
        flash("Преподаватель не найден.")
        return redirect(url_for('cabinet'))

    teacher_id = teacher['id']
    cur.execute("""
        SELECT s.date, s.time_slot, g.group_name, sub.name AS subject_name
        FROM schedules s
        JOIN grooups g ON s.group_id = g.id
        JOIN subjects sub ON s.subject_id = sub.id
        WHERE s.teacher_id = %s AND s.date = CURDATE()
    """, (teacher_id,))
    schedule = cur.fetchall()
    cur.close()
    return render_template('teacher_schedule.html', schedule=schedule)


# Добавление записи в расписание
@app.route('/add_schedule', methods=['GET', 'POST'])
@role_required('admin')
def add_schedule():
    if request.method == 'POST':
        
        group_id = request.form['group_id']
        subject_id = request.form['subject_id']
        teacher_id = request.form['teacher_id']
        date = request.form['date']
        time_slot = request.form['time_slot']

        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO schedules (group_id, subject_id, teacher_id, date, time_slot)
            VALUES (%s, %s, %s, %s, %s)
        """, (group_id, subject_id, teacher_id, date, time_slot))
        mysql.connection.commit()
        cur.close()
        log_user_action(session.get('user_id'), "Добавлена новая запись", request.method)
        return redirect(url_for('view_schedule'))
    
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM grooups")
    groups = cur.fetchall()
    cur.execute("SELECT * FROM subjects")
    subjects = cur.fetchall()
    cur.execute("SELECT * FROM teachers")
    teachers = cur.fetchall()
    cur.close()
    return render_template('add_schedule.html', groups=groups, subjects=subjects, teachers=teachers)

# Удаление записи из расписания
@app.route('/delete_schedule/<int:id>')
@role_required('admin')
def delete_schedule(id):

    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM schedules WHERE id = %s", (id,))
    mysql.connection.commit()
    cur.close()
    log_user_action(session.get('user_id'), f"Удалено расписание с ID {id}", request.method)
    return redirect(url_for('view_schedule'))

# Поиск расписания
@app.route('/search_schedule', methods=['GET', 'POST'])
def search_schedule():
    results = []
    if request.method == 'POST':
        search_query = request.form['search_query']
        cur = mysql.connection.cursor()
        query = """
            SELECT s.id, g.group_name, sub.name AS subject_name, t.full_name AS teacher_name, 
                   s.date, s.time_slot
            FROM schedules s
            JOIN grooups g ON s.group_id = g.id
            JOIN subjects sub ON s.subject_id = sub.id
            JOIN teachers t ON s.teacher_id = t.id
            WHERE g.group_name LIKE %s OR sub.name LIKE %s OR t.full_name LIKE %s
        """
        cur.execute(query, (f"%{search_query}%", f"%{search_query}%", f"%{search_query}%"))
        results = cur.fetchall()
        cur.close()
        log_user_action(session.get('user_id'), "Произведён поиск", request.method)
    return render_template('search_schedule.html', results=results)

# Сортировка расписания
@app.route('/sort_schedule', methods=['GET'])
def sort_schedule():
    sort_by = request.args.get('sort_by', 'date')  # Поле для сортировки, по умолчанию "date"
    if sort_by not in ['group_name', 'subject_name', 'teacher_name', 'date', 'time_slot']:
        sort_by = 'date'  # Если указанное поле недопустимо, сортируем по дате

    cur = mysql.connection.cursor()
    query = f"""
        SELECT s.id, g.group_name, sub.name AS subject_name, t.full_name AS teacher_name, 
               s.date, s.time_slot
        FROM schedules s
        JOIN grooups g ON s.group_id = g.id
        JOIN subjects sub ON s.subject_id = sub.id
        JOIN teachers t ON s.teacher_id = t.id
        ORDER BY {sort_by}
    """
    cur.execute(query)
    sorted_schedule = cur.fetchall()
    cur.close()

    # Определяем роль пользователя (например, admin, teacher, student)
    role = session.get('role', 'guest')  # 'guest' по умолчанию
    log_user_action(session.get('user_id'), "Произошла сортировка", request.method)
    return render_template('sort_schedule.html', schedule=sorted_schedule, role=role)


# Замена записи в расписании
@app.route('/substitute/<int:schedule_id>', methods=['GET', 'POST'])
@role_required('admin')
def substitute(schedule_id):
    if request.method == 'POST':
        # Получение данных из формы
        new_subject_id = request.form['new_subject_id']
        new_teacher_id = request.form['new_teacher_id']
        replacement_date = request.form['replacement_date']
        is_half_lesson = 'is_half_lesson' in request.form  # Булевое значение для половины занятия

        # Обновление или добавление замены в базе данных
        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO replacements (schedule_id, new_subject_id, new_teacher_id, replacement_date, is_half_lesson)
            VALUES (%s, %s, %s, %s, %s)
        """, (schedule_id, new_subject_id, new_teacher_id, replacement_date, is_half_lesson))
        mysql.connection.commit()
        cur.close()
        log_user_action(session.get('user_id'), f"Замена в расписании ID {schedule_id}", request.method)
        flash('Замена успешно добавлена.')
        return redirect(url_for('view_schedule'))

    # Получение данных для отображения текущей записи и выбора замены
    cur = mysql.connection.cursor()
    cur.execute("""
        SELECT s.id, g.group_name, sub.name AS subject_name, t.full_name AS teacher_name, s.date, s.time_slot
        FROM schedules s
        JOIN grooups g ON s.group_id = g.id
        JOIN subjects sub ON s.subject_id = sub.id
        JOIN teachers t ON s.teacher_id = t.id
        WHERE s.id = %s
    """, (schedule_id,))
    schedule = cur.fetchone()

    # Получение списка доступных предметов и преподавателей
    cur.execute("SELECT * FROM subjects")
    subjects = cur.fetchall()
    cur.execute("SELECT * FROM teachers")
    teachers = cur.fetchall()
    cur.close()

    return render_template('substitute.html', schedule=schedule, subjects=subjects, teachers=teachers)



#=======Автогенерация========#
@app.route('/auto_generate')
@role_required('admin')
def auto_generate():
    cur = mysql.connection.cursor()

    # Получаем активный семестр
    cur.execute("SELECT * FROM semester_settings ORDER BY id DESC LIMIT 1")
    semester = cur.fetchone()
    if not semester:
        flash('Сначала задайте параметры семестра.')
        return redirect(url_for('semester_settings'))

    start_date = semester['start_date']
    end_date = semester['end_date']

    # Получаем праздники
    cur.execute("SELECT holiday_date FROM holidays")
    holidays = set(row['holiday_date'] for row in cur.fetchall())

    # Получаем расписание звонков
    cur.execute("SELECT pair_number FROM call_schedule")
    pairs = [row['pair_number'] for row in cur.fetchall()]
    pairs.sort()

    # Получаем группы
    cur.execute("SELECT * FROM grooups")
    groups = cur.fetchall()

    # Получаем связи преподаватель-предмет
    cur.execute("""
        SELECT ts.teacher_id, ts.subject_id, s.hours_per_week
        FROM teacher_subjects ts
        JOIN subjects s ON s.id = ts.subject_id
    """)
    teacher_subjects = cur.fetchall()

    # Очищаем расписание перед генерацией
    cur.execute("DELETE FROM replacements")
    cur.execute("DELETE FROM schedules")
    mysql.connection.commit()

    # Вспомогательная функция: занят ли преподаватель
    def is_teacher_busy(teacher_id, date, time_slot):
        cur.execute("""
            SELECT COUNT(*) AS cnt FROM schedules
            WHERE teacher_id = %s AND date = %s AND time_slot = %s
        """, (teacher_id, date, time_slot))
        return cur.fetchone()['cnt'] > 0

    # Вспомогательная функция: сколько пар у группы в день
    def group_pairs_count(group_id, date):
        cur.execute("""
            SELECT COUNT(*) AS cnt FROM schedules
            WHERE group_id = %s AND date = %s
        """, (group_id, date))
        return cur.fetchone()['cnt']

    # Начинаем генерацию по дням
    current_date = start_date

    while current_date <= end_date:
        if current_date.weekday() >= 5 or current_date in holidays:
            current_date += timedelta(days=1)
            continue  # Пропуск выходных и праздников

        for group in groups:
            group_id = group['id']

            for ts in teacher_subjects:
                teacher_id = ts['teacher_id']
                subject_id = ts['subject_id']
                hours_per_week = ts['hours_per_week']

                # Проход по всем парам дня
                for pair in pairs:
                    if group_pairs_count(group_id, current_date) >= 6:  # макс 6 пар в день
                        break

                    if is_teacher_busy(teacher_id, current_date, f"{pair} пара"):
                        continue  # если преподаватель занят

                    # Вставляем в расписание
                    cur.execute("""
                        INSERT INTO schedules (group_id, subject_id, teacher_id, date, time_slot)
                        VALUES (%s, %s, %s, %s, %s)
                    """, (group_id, subject_id, teacher_id, current_date, f"{pair} пара"))

        current_date += timedelta(days=1)

    mysql.connection.commit()
    cur.close()
    flash('Расписание успешно сгенерировано без конфликтов.')
    return redirect(url_for('view_schedule'))


# -------------------------------
# Управление группами (только admin)
# -------------------------------
@app.route('/groups')
@role_required('admin')
def manage_groups():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM grooups")
    groups = cur.fetchall()
    cur.close()
    return render_template('groups.html', groups=groups)

@app.route('/groups/add', methods=['POST'])
@role_required('admin')
def add_group():
    group_name = request.form['group_name']
    cur = mysql.connection.cursor()
    cur.execute("INSERT INTO grooups (group_name) VALUES (%s)", (group_name,))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Добавлена группа: {group_name}", request.method)
    return redirect(url_for('manage_groups'))

@app.route('/groups/delete/<int:id>')
@role_required('admin')
def delete_group(id):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM grooups WHERE id = %s", (id,))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Удалена группа ID: {id}", request.method)
    return redirect(url_for('manage_groups'))

# -------------------------------
# Управление дисциплинами (только admin)
# -------------------------------
@app.route('/subjects')
@role_required('admin')
def manage_subjects():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM subjects")
    subjects = cur.fetchall()
    cur.close()
    return render_template('subjects.html', subjects=subjects)

@app.route('/subjects/add', methods=['POST'])
@role_required('admin')
def add_subject():
    name = request.form['name']
    total_hours = request.form['total_hours']
    hours_per_week = request.form['hours_per_week']
    cur = mysql.connection.cursor()
    cur.execute("""
        INSERT INTO subjects (name, total_hours, hours_per_week)
        VALUES (%s, %s, %s)
    """, (name, total_hours, hours_per_week))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Добавлена дисциплина: {name}", request.method)
    return redirect(url_for('manage_subjects'))

@app.route('/subjects/delete/<int:id>')
@role_required('admin')
def delete_subject(id):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM subjects WHERE id = %s", (id,))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Удалена дисциплина ID: {id}", request.method)
    return redirect(url_for('manage_subjects'))

# -------------------------------
# Управление преподавателями (только admin)
# -------------------------------
@app.route('/teachers')
@role_required('admin')
def manage_teachers():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM teachers")
    teachers = cur.fetchall()
    cur.close()
    return render_template('teachers.html', teachers=teachers)

@app.route('/teachers/add', methods=['POST'])
@role_required('admin')
def add_teacher():
    full_name = request.form['full_name']
    department = request.form['department']
    cur = mysql.connection.cursor()
    cur.execute("INSERT INTO teachers (full_name, department) VALUES (%s, %s)", (full_name, department))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Добавлен преподаватель: {full_name}", request.method)
    return redirect(url_for('manage_teachers'))

@app.route('/teachers/delete/<int:id>')
@role_required('admin')
def delete_teacher(id):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM teachers WHERE id = %s", (id,))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Удалён преподаватель ID: {id}", request.method)
    return redirect(url_for('manage_teachers'))

# -------------------------------
# Управление аудиториями (только admin)
# -------------------------------
@app.route('/auditoriums')
@role_required('admin')
def manage_auditoriums():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM auditoriums")
    auditoriums = cur.fetchall()
    cur.close()
    return render_template('auditoriums.html', auditoriums=auditoriums)

@app.route('/auditoriums/add', methods=['POST'])
@role_required('admin')
def add_auditorium():
    name = request.form['name']
    capacity = request.form['capacity']
    building = request.form['building']
    cur = mysql.connection.cursor()
    cur.execute("INSERT INTO auditoriums (name, capacity, building) VALUES (%s, %s, %s)", (name, capacity, building))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Добавлена аудитория: {name}", request.method)
    return redirect(url_for('manage_auditoriums'))

@app.route('/auditoriums/delete/<int:id>')
@role_required('admin')
def delete_auditorium(id):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM auditoriums WHERE id = %s", (id,))
    mysql.connection.commit()
    cur.close()
    log_user_action(session['user_id'], f"Удалена аудитория ID: {id}", request.method)
    return redirect(url_for('manage_auditoriums'))


# -------------------------------
# Параметры семестра (только admin)
# -------------------------------
@app.route('/semester', methods=['GET', 'POST'])
@role_required('admin')
def semester_settings():
    cur = mysql.connection.cursor()

    if request.method == 'POST':
        semester_name = request.form['semester_name']
        start_date = request.form['start_date']
        end_date = request.form['end_date']
        weeks_count = request.form['weeks_count']

        cur.execute("""
            INSERT INTO semester_settings (semester_name, start_date, end_date, weeks_count)
            VALUES (%s, %s, %s, %s)
        """, (semester_name, start_date, end_date, weeks_count))
        # Очищаем расписание
        cur.execute("DELETE FROM replacements")
        cur.execute("DELETE FROM schedules")
        
        mysql.connection.commit()
        log_user_action(session['user_id'], f"Изменены параметры семестра: {semester_name}", request.method)

    cur.execute("SELECT * FROM semester_settings ORDER BY id DESC LIMIT 1")
    semester = cur.fetchone()
    cur.close()
    return render_template('settings.html', semester=semester)



# -------------------------------
# График звонков (только admin)
# -------------------------------
@app.route('/call_schedule', methods=['GET', 'POST'])
@role_required('admin')
def call_schedule():
    cur = mysql.connection.cursor()

    if request.method == 'POST':
        pair_number = request.form['pair_number']
        start_time = request.form['start_time']
        end_time = request.form['end_time']

        cur.execute("""
            INSERT INTO call_schedule (pair_number, start_time, end_time)
            VALUES (%s, %s, %s)
        """, (pair_number, start_time, end_time))
        mysql.connection.commit()
        log_user_action(session['user_id'], "Добавлен звонок", request.method)

    cur.execute("SELECT * FROM call_schedule ORDER BY pair_number")
    schedule = cur.fetchall()
    cur.close()
    return render_template('call_schedule.html', schedule=schedule)


# -------------------------------
# Управление праздниками (только admin)
# -------------------------------
@app.route('/holidays', methods=['GET', 'POST'])
@role_required('admin')
def holidays():
    cur = mysql.connection.cursor()

    if request.method == 'POST':
        holiday_date = request.form['holiday_date']
        description = request.form['description']
        cur.execute("INSERT INTO holidays (holiday_date, description) VALUES (%s, %s)", (holiday_date, description))
        mysql.connection.commit()
        log_user_action(session['user_id'], f"Добавлен выходной: {holiday_date}", request.method)

    cur.execute("SELECT * FROM holidays ORDER BY holiday_date")
    holidays = cur.fetchall()
    cur.close()
    return render_template('holidays.html', holidays=holidays)



# -------------------------------
# Разовые мероприятия (только admin)
# -------------------------------
@app.route('/events', methods=['GET', 'POST'])
@role_required('admin')
def events():
    cur = mysql.connection.cursor()

    if request.method == 'POST':
        title = request.form['event_title']
        description = request.form['description']
        date = request.form['event_date']
        time_slot = request.form['time_slot']
        group_id = request.form['group_id']
        auditorium_id = request.form['auditorium_id']

        cur.execute("""
            INSERT INTO events (event_title, description, event_date, time_slot, group_id, auditorium_id)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (title, description, date, time_slot, group_id, auditorium_id))
        mysql.connection.commit()
        log_user_action(session['user_id'], f"Добавлено мероприятие: {title}", request.method)

    cur.execute("SELECT * FROM events ORDER BY event_date")
    events = cur.fetchall()
    cur.execute("SELECT * FROM grooups")
    groups = cur.fetchall()
    cur.execute("SELECT * FROM auditoriums")
    auditoriums = cur.fetchall()
    cur.close()

    return render_template('events.html', events=events, groups=groups, auditoriums=auditoriums)

@app.route('/export_workload_report')
@role_required('admin')
def export_workload_report():
    cur = mysql.connection.cursor()

    # Получаем расписание за текущую неделю
    cur.execute("""
        SELECT t.full_name AS teacher_name, COUNT(*) AS total_lessons
        FROM schedules s
        JOIN teachers t ON s.teacher_id = t.id
        WHERE WEEK(s.date) = WEEK(CURDATE())
        GROUP BY t.full_name
    """)
    result = cur.fetchall()

    # Преобразуем к списку словарей (на случай, если вернулось как tuple)
    data = []
    for row in result:
        if isinstance(row, dict):
            data.append(row)
        else:
            data.append({'teacher_name': row[0], 'total_lessons': row[1]})

    # Получаем всех преподавателей
    cur.execute("SELECT full_name FROM teachers")
    all_teachers = [r[0] if not isinstance(r, dict) else r['full_name'] for r in cur.fetchall()]
    cur.close()

    # Добавляем отсутствующих преподавателей с 0 пар
    existing_names = [row['teacher_name'] for row in data]
    for t in all_teachers:
        if t not in existing_names:
            data.append({'teacher_name': t, 'total_lessons': 0})

    # Формируем Excel


    wb = Workbook()
    ws = wb.active
    ws.title = "Нагрузка преподавателей"

    ws.append(["Преподаватель", "Количество пар", "Комментарий"])
    for row in data:
        name = row['teacher_name']
        count = row['total_lessons']
        comment = "Этот преподаватель не может работать на этой неделе" if count == 0 else \
                  f"На этой неделе этот преподаватель может проработать {count * 1.5:.1f} часов"
        ws.append([name, count, comment])

    for cell in ws[1]:
        cell.font = Font(bold=True)

    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    log_user_action(session['user_id'], "Экспорт отчета по нагрузке преподавателей", "GET")
    return send_file(file_stream, as_attachment=True, download_name="workload_report.xlsx")



@app.route('/export_full_schedule_analysis')
@role_required('admin')
def export_full_schedule_analysis():
    cur = mysql.connection.cursor()

    # Получаем расписание
    cur.execute("""
        SELECT s.id, g.group_name, t.full_name AS teacher_name, s.date, s.time_slot
        FROM schedules s
        JOIN grooups g ON s.group_id = g.id
        JOIN teachers t ON s.teacher_id = t.id
        ORDER BY s.date, s.time_slot
    """)
    schedule = cur.fetchall()

    # Получаем всех преподавателей
    cur.execute("SELECT full_name FROM teachers")
    all_teachers = [row['full_name'] for row in cur.fetchall()]
    cur.close()

    # Подготовка данных


    group_gaps = defaultdict(lambda: defaultdict(list))
    teacher_conflicts = defaultdict(lambda: defaultdict(list))
    teacher_workload = defaultdict(int)

    for row in schedule:
        group_gaps[row['group_name']][row['date']].append(row['time_slot'])
        teacher_conflicts[row['teacher_name']][row['date']].append(row['time_slot'])
        teacher_workload[row['teacher_name']] += 1

    # Анализ форточек
    windows = []
    for group, days in group_gaps.items():
        for day, slots in days.items():
            pairs = sorted(int(s.split()[0]) for s in slots if s[0].isdigit())
            for i in range(1, len(pairs)):
                if pairs[i] - pairs[i - 1] > 1:
                    windows.append({
                        "group": group,
                        "date": day,
                        "comment": f"Между {pairs[i - 1]} и {pairs[i]} парой"
                    })

    # Анализ конфликтов
    conflicts = []
    for teacher, days in teacher_conflicts.items():
        for day, slots in days.items():
            duplicates = set([x for x in slots if slots.count(x) > 1])
            for conflict in duplicates:
                conflicts.append({
                    "teacher": teacher,
                    "date": day,
                    "pair": conflict,
                    "comment": "Совпадение пар"
                })

    # Объединяем нагрузку (с учётом преподавателей без занятий)
    all_data = []
    for name in all_teachers:
        lessons = teacher_workload.get(name, 0)
        if lessons == 0:
            comment = "Этот преподаватель не может работать на этой неделе"
        else:
            comment = f"На этой неделе этот преподаватель может проработать {lessons * 1.5:.1f} часов"
        all_data.append({"teacher": name, "lessons": lessons, "comment": comment})

    # Excel генерация
    wb = Workbook()

    # Форточки
    ws1 = wb.active
    ws1.title = "Форточки"
    ws1.append(["Группа", "Дата", "Комментарий"])
    for row in windows:
        ws1.append([row['group'], row['date'], row['comment']])
    for cell in ws1[1]: cell.font = Font(bold=True)

    # Конфликты
    ws2 = wb.create_sheet("Конфликты")
    ws2.append(["Преподаватель", "Дата", "Пара", "Комментарий"])
    for row in conflicts:
        ws2.append([row['teacher'], row['date'], row['pair'], row['comment']])
    for cell in ws2[1]: cell.font = Font(bold=True)

    # Нагрузка
    ws3 = wb.create_sheet("Нагрузка")
    ws3.append(["Преподаватель", "Количество пар", "Комментарий"])
    for row in all_data:
        ws3.append([row['teacher'], row['lessons'], row['comment']])
    for cell in ws3[1]: cell.font = Font(bold=True)

    # Отправка
    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    log_user_action(session['user_id'], "Экспорт полного анализа расписания", "GET")
    return send_file(file_stream, as_attachment=True, download_name="full_schedule_analysis.xlsx")


@app.route('/export_schedule_all')
@role_required('admin')
def export_schedule_all():
    cur = mysql.connection.cursor()
    cur.execute("""
        SELECT s.date, g.group_name, sub.name AS subject_name, t.full_name AS teacher_name, s.time_slot
        FROM schedules s
        JOIN grooups g ON s.group_id = g.id
        JOIN subjects sub ON s.subject_id = sub.id
        JOIN teachers t ON s.teacher_id = t.id
        ORDER BY s.date, s.time_slot, g.group_name
    """)
    rows = cur.fetchall()
    cur.close()

    # === Excel ===
    wb = Workbook()
    ws = wb.active
    ws.title = "Расписание"
    ws.append(["Дата", "День недели", "Группа", "Пара", "Предмет", "Преподаватель"])
    
    for row in rows:
        date_str = row['date']
        weekday = datetime.strptime(str(date_str), "%Y-%m-%d").strftime("%A")
        ws.append([date_str, weekday, row['group_name'], row['time_slot'], row['subject_name'], row['teacher_name']])

    excel_stream = BytesIO()
    wb.save(excel_stream)
    excel_stream.seek(0)

    # === Word ===
    doc = Document()
    doc.add_heading("Расписание на неделю", 0)

    current_date = ""
    for row in rows:
        date_str = row['date']
        weekday = datetime.strptime(str(date_str), "%Y-%m-%d").strftime("%A")

        if current_date != date_str:
            doc.add_heading(f"{date_str} ({weekday})", level=1)
            current_date = date_str

        doc.add_paragraph(
            f"Группа: {row['group_name']}, Пара: {row['time_slot']}, "
            f"{row['subject_name']} — {row['teacher_name']}"
        )

    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)

    # === ZIP архив ===
    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, 'w') as zipf:
        zipf.writestr("schedule.xlsx", excel_stream.getvalue())
        zipf.writestr("schedule.docx", word_stream.getvalue())

    zip_buffer.seek(0)
    log_user_action(session['user_id'], "Экспорт расписания в Excel и Word", "GET")

    return send_file(zip_buffer, as_attachment=True, download_name="schedule_export.zip")


if __name__ == '__main__':
    app.run(debug=True)
   

